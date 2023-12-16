from typing import Counter
from mailmerge import MailMerge
from docx.shared import Inches
from docx import Document
import numpy as np
import matplotlib
matplotlib.use('Agg')  # 设置 Matplotlib 使用非交互式后端
import matplotlib.pyplot as plt
from matplotlib.ticker import FixedLocator
from matplotlib.ticker import FuncFormatter
import datetime
import os
import io

def create_peer_relationship(fingers):
    # 建立 peer 關係
    fingers["R1"].set_peer(fingers["L1"])
    fingers["R2"].set_peer(fingers["L2"])
    fingers["R3"].set_peer(fingers["L3"])
    fingers["R4"].set_peer(fingers["L4"])
    fingers["R5"].set_peer(fingers["L5"])

def calculate_tmp_ranking(fingers):
    sorted_fingers = sorted(fingers, key=lambda finger: (
        max(finger.left_value, finger.right_value), 
        min(finger.left_value, finger.right_value), 
        finger.pattern_weight,
        finger.hand_weight
    ), reverse=True)

    current_ranking = 1
    last_finger_key = None
    for i, finger in enumerate(sorted_fingers, start=1):
        finger_key = (
            max(finger.left_value, finger.right_value), 
            min(finger.left_value, finger.right_value), 
            finger.pattern_weight,
            finger.hand_weight
        )
        
        # 檢查是否與上一個手指具有相同的排序鍵
        if finger_key != last_finger_key:
            current_ranking = i

        finger.tmp_ranking = current_ranking
        last_finger_key = finger_key

def create_additional_info(fingers):
    W_ranking_list = ["Wp", "We", "Wt", "Ws", "Wc"]
    W_ranking = [5, 2, 4, 3, 1]
    U_A_ranking_list = ["U", "Au", "Ar", "As", "At", "R"]
    U_A_ranking = [5, 4, 3, 2, 1]

    for name, finger in fingers.items():
        finger_num = name[1:]

        if(finger.pattern in W_ranking_list):
            finger.hand_ranking = W_ranking[int(finger_num) - 1]
        elif(finger.pattern in U_A_ranking_list):
            finger.hand_ranking = U_A_ranking[int(finger_num) - 1]

def sort_fingers(fingers):

    sorted_fingers = sorted(fingers, key=lambda finger: (
        finger.tmp_ranking,
        finger.peer_tmp_ranking
    ))

    last_finger_key = None
    group_fingers = []
    for i, finger in enumerate(sorted_fingers, start=1):
        finger_key = (
            finger.tmp_ranking,
            finger.peer_tmp_ranking
        )
        # 檢查是否與上一個手指具有相同的排序鍵
        if i == 1 or finger_key == last_finger_key:
            group_fingers.append(finger)
        else:
            sorted_group = sorted(group_fingers, key=lambda finger: (
                finger.peer_ranking,
                finger.hand_ranking
            ))

            for rank, _finger in enumerate(sorted_group, start=0):
                _finger.ranking = rank + i - len(sorted_group)

            # 跟新 group_fingers
            group_fingers = []
            group_fingers.append(finger)

        last_finger_key = finger_key
    
    # 處理剩餘的排名
    sorted_group = sorted(group_fingers, key=lambda finger: (
        finger.peer_ranking,
        finger.hand_ranking
    ))

    for rank, _finger in enumerate(sorted_group, start=0):
        # 上面的 for 跳出之後 i 是 10 不是 11
        _finger.ranking = rank + (i+1) - len(sorted_group)

def calculate_percentage(fingers, total_value):
    for finger in fingers.values():
        finger.percentage = round(finger.value / total_value * 100.0, 1)

def calculate_pattern_percentage(fingers):
    pattern_map_4 = {
        "We": "W",
        "Wt": "W",
        "Ws": "W",
        "Wp": "W",
        "Wc": "W",
        "U" : "U",
        "R" : "R",
        "As": "A",
        "Au": "A",
        "Ar": "A",
        "At": "A",
    }
    pattern_percentage = Counter()
    for finger in fingers.values():
        pattern_percentage[pattern_map_4[finger.pattern]] += 1

    return pattern_percentage

def get_eight_wisdoms(fingers):
    eight_wisdoms = {
        "interpersonal" : fingers["L1"].value,
        "introspection" : fingers["R1"].value,
        "image" : (fingers["L2"].value + fingers["L5"].value) / 2.0,
        "math" : fingers["R2"].value,
        "rhythm" : (fingers["L3"].value + fingers["R3"].value) / 2.0,
        "music" : fingers["L4"].value,
        "language" : fingers["R4"].value,
        "nature" : fingers["R5"].value
    }

    total_value = sum(eight_wisdoms.values())
    E = {key: (value / total_value) * 100 for key, value in eight_wisdoms.items()}
    
    return E

def create_language(fingers):
    forces = {}
    name = ["feeling", "understanding", "writing", "speaking", "reading"]

    for i in range(1, 6):        
        force_value = fingers[f"L{i}"].value + fingers[f"R{i}"].value
        force_tuple = (force_value, i)
        
        forces[name[i-1]] = force_tuple

    return forces

def create_music(fingers):
    forces = {}
    name = ["feeling", "creations", "preformance", "sensitivity", "appreciation"]

    for i in range(1, 6):        
        force_value = fingers[f"L{i}"].value + fingers[f"R{i}"].value
        force_tuple = (force_value, i)
        
        forces[name[i-1]] = force_tuple

    return forces

def create_logic(fingers):
    forces = {}
    name = ["experiment", "speculation", "calculation", "deduction", "induction"]

    for i in range(1, 6):        
        force_value = fingers[f"L{i}"].value + fingers[f"R{i}"].value
        force_tuple = (force_value, i)
        
        forces[name[i-1]] = force_tuple

    return forces

def create_limbs(fingers, pattern_map):
    forces = {}
    name = ["reaction", "rhythm", "control", "sensation", "balance"]
    finger_number = [1, 3, 3, 3, 5]
    finger_weight = [1, 3, 3, 3, 5]

    for i, num in enumerate(finger_number, start=0):
        ranking = 0
        pattern_score = 0
        control_pattern_scores = {
            'A': 4,   # Represents As, Au, Ar, At
            'Ws': 3,  # Represents We, Wt, Ws
            'Wc': 2,  # Represents Wc, Wp
            'U': 1, 
            'R': 1
        }
        sensation_pattern_scores = {
            'U': 4,
            'R': 4,
            'Wc': 3,  # Represents Wc, Wp
            'Ws': 2,  # Represents We, Wt, Ws
            'A': 1    # Represents As, Au, Ar, At
        }
        if(name[i] == "rhythm"):
            ranking = fingers["L3"].ranking
            pattern_score = 0
        elif(name[i] == "control" or name[i] == "sensation"):
            ranking = fingers["R3"].ranking
            if(name[i] == "control"):
                pattern_score = control_pattern_scores[pattern_map[fingers["R3"].pattern]]
            else:
                pattern_score = sensation_pattern_scores[pattern_map[fingers["R3"].pattern]]
        else:
            ranking = 0
            pattern_score = 0

        force_value = fingers[f"L{num}"].value + fingers[f"R{num}"].value
        #為了把 ranking 變成分數，所以用 (11-ranking) 來當成分數。
        force_tuple = (force_value, finger_weight[i], 11 - ranking, pattern_score)
        
        forces[name[i]] = force_tuple

    return forces

def create_vision(fingers, pattern_map):
    forces = {}
    name = ["space", "assemble", "drawing", "image", "color"]
    finger_number = [2, 3, 3, 5, 5]
    finger_weight = [2, 3, 3, 5, 5]

    for i, num in enumerate(finger_number, start=0):
        ranking = 0
        pattern_score = 0
        image_pattern_scores = {
            'U': 4,
            'R': 4,
            'Wc': 3,  # This will represent Wc, Wp
            'Ws': 2,  # This will represent We, Wt, Ws
            'A': 1    # This will represent As, Au, Ar, At
        }
        color_pattern_scores = {
            'Ws': 4,  # This will represent We, Wt, Ws
            'U': 3,
            'R': 3,
            'A': 2,   # This will represent As, Au, Ar, At
            'Wc': 1   # This will represent Wc, Wp
        }
        if(name[i] == "assemble" or name[i] == "drawing"):
            if(name[i] == "assemble"):
                ranking = fingers["L2"].ranking
            else:
                ranking = fingers["L5"].ranking
            pattern_score = 0

        elif(name[i] == "image" or name[i] == "color"):
            ranking = 0
            if(name[i] == "image"):
                pattern_score = image_pattern_scores[pattern_map[fingers["L5"].pattern]]
            else:
                pattern_score = color_pattern_scores[pattern_map[fingers["L5"].pattern]]
        else:
            ranking = 0
            pattern_score = 0

        force_value = fingers[f"L{num}"].value + fingers[f"R{num}"].value
        #為了把 ranking 變成分數，所以用 (11-ranking) 來當成分數。
        force_tuple = (force_value, finger_weight[i], 11 - ranking, pattern_score)
        
        forces[name[i]] = force_tuple

    return forces

def create_interpersonal(fingers, pattern_map):
    forces = {}
    name = ["leadership", "empathy", "understanding", "collaboration", "communication"]
    finger_number = [1, 1, 2, 3, 4]
    finger_weight = [1, 1, 2, 3, 4]

    for i, num in enumerate(finger_number, start=0):
        pattern_score = 0
        leadership_pattern_scores = {
            'Ws': 4,  # This will represent We, Wt, Ws
            'A': 3,   # This will represent As, Au, Ar, At
            'Wc': 2,  # This will represent Wc, Wp
            'U': 1, 
            'R': 1
        }
        empathy_pattern_scores = {
            'U': 4,
            'R': 4,
            'Wc': 3,  # This will represent Wc, Wp
            'Ws': 2,  # This will represent We, Wt, Ws
            'A': 1    # This will represent As, Au, Ar, At
        }
        if(name[i] == "leadership"):
            pattern_score = leadership_pattern_scores[pattern_map[fingers["L1"].pattern]]
        elif(name[i] == "empathy"):
            pattern_score = empathy_pattern_scores[pattern_map[fingers["L1"].pattern]]
        else:
            pattern_score = 0
        
        force_value = fingers[f"L{num}"].value + fingers[f"R{num}"].value
        force_tuple = (force_value, finger_weight[i], pattern_score)
        
        forces[name[i]] = force_tuple

    return forces

def create_nature(fingers):
    forces = {}
    name = ["concern", "categorization", "exploration", "observation", "identification"]
    finger_number = [1, 2, 3, 5, 5]
    finger_weight = [1, 2, 3, 5, 5]

    for i, num in enumerate(finger_number, start=0):
        ranking = 0
        if(name[i] == "observation"):
            ranking = fingers["L5"].ranking
        elif(name[i] == "identification"):
            ranking = fingers["R5"].ranking
        else:
            ranking = 0
        
        force_value = fingers[f"L{num}"].value + fingers[f"R{num}"].value
        force_tuple = (force_value, finger_weight[i], ranking)
        
        forces[name[i]] = force_tuple

    return forces

def create_introspection(fingers, pattern_map):
    forces = {}
    name = ["discipline", "compassion", "reflection", "empathy", "insight"]
    finger_number = [1, 1, 1, 1, 5]
    finger_weight = [1, 1, 1, 1, 5]

    for i, num in enumerate(finger_number, start=0):
        pattern_score = 0
        discipline_pattern_scores = {
            'A': 4, 'Ws': 3, 'Wc': 2, 'U': 1, 'R': 1
        }
        compassion_pattern_scores = {
            'Wc': 4, 'A': 3, 'U': 2, 'R': 2, 'Ws': 1
        }
        reflection_pattern_scores = {
            'Ws': 4, 'U': 3, 'R': 3, 'A': 2, 'Wc': 1
        }
        empathy_pattern_scores = {
            'U': 4, 'R': 4, 'Wc': 3, 'Ws': 2, 'A': 1
        }
        if(name[i] == "discipline"):
            pattern_score = discipline_pattern_scores[pattern_map[fingers["R1"].pattern]]
        elif(name[i] == "compassion"):
            pattern_score = compassion_pattern_scores[pattern_map[fingers["R1"].pattern]]
        elif(name[i] == "reflection"):
            pattern_score = reflection_pattern_scores[pattern_map[fingers["R1"].pattern]]
        elif(name[i] == "empathy"):
            pattern_score = empathy_pattern_scores[pattern_map[fingers["R1"].pattern]]
        else:
            pattern_score = 0
        
        force_value = fingers[f"L{num}"].value + fingers[f"R{num}"].value
        force_tuple = (force_value, finger_weight[i], pattern_score)
        
        forces[name[i]] = force_tuple

    return forces

def f5_analysis(fingers, pattern_map):
    #Create tuples for five forces analysis
    dict_datas = {
        "language" : create_language(fingers),
        "music" : create_music(fingers),
        "logic" : create_logic(fingers),
        "limbs" : create_limbs(fingers, pattern_map),
        "vision" : create_vision(fingers, pattern_map),
        "interpersonal" : create_interpersonal(fingers, pattern_map),
        "nature" : create_nature(fingers),
        "introspection" : create_introspection(fingers, pattern_map)
    }

    # 定義一個自定義的比較函數
    def custom_sort(item):
        return item[1]  # 使用tuple的第二個元素作為排序依據

    f5_rankings = {}
    for force_name, dict_data in dict_datas.items():
        # 使用sorted函數對字典中的項目進行排序，並將其轉換為列表
        sorted_data = sorted(dict_data.items(), key=custom_sort, reverse=True)
        
        rankings = {}
        for rank, (key, value) in enumerate(sorted_data, start=1):
            rankings[key] = rank
        
        f5_rankings[force_name] = rankings

    return f5_rankings

def get_information(prefix, key, fingers=None, pattern_map=None):
    if fingers and pattern_map:
        try:
            file_name = pattern_map[fingers[key].pattern]
        except KeyError:
            raise ValueError(f'{fingers[key].pattern} 不在已知的文字庫中')
    else:
        file_name = key

    dirname = os.path.dirname(os.path.dirname(os.path.realpath(__file__)))
    file_path = os.path.join(dirname, 'database', f'{prefix}{file_name}.txt')

    try:
        with open(file_path, 'r', encoding="utf-8") as file:
            content = ''.join(file)
    except FileNotFoundError:
        raise FileNotFoundError(f"檔案 '{file_path}' 不存在!")

    return content

def get_data(fingers, pattern_map):

    total_value = sum([finger.value for finger in fingers.values()])
    calculate_percentage(fingers, total_value)

    pattern_percentage = calculate_pattern_percentage(fingers)

    E = get_eight_wisdoms(fingers)

    # 獲取排名的前置作業（peer -> tmp_ranking -> add_info -> sort）
    create_peer_relationship(fingers)
    calculate_tmp_ranking(fingers.values())
    create_additional_info(fingers)
    sort_fingers(fingers.values())
    top3_fingers = {finger.ranking: finger for finger in fingers.values() if finger.ranking in [1, 2, 3]}

    f5_rankings = f5_analysis(fingers, pattern_map)
    # debug
    # for force_name, rankings in f5_rankings.items():
    #     print(force_name)
    #     for feature, rank in rankings.items():
    #         print(feature, rank)
    #     print("----")

    career_map = {
        "L1": {
            "pattern": {
                "Ws" : "精準型", 
                "Wc" : "多元型", 
                "U" : "服務型",
                "R" : "服務型", 
                "A" : "務實型"
            },
            "protential": "人際潛能"
        },
        "L2": {
            "pattern": {
                "Ws" : "精準型", 
                "Wc" : "多元型", 
                "U" : "創意型",
                "R" : "創意型", 
                "A" : "務實型"
            },
            "protential": "空間潛能"
        },
        "L3": {
            "pattern": {
                "Ws" : "力態型", 
                "Wc" : "多元型", 
                "U" : "美感型",
                "R" : "美感型", 
                "A" : "慣性型"
            },
            "protential": "律動潛能"
        },
        "L4": {
            "pattern": {
                "Ws" : "精準型", 
                "Wc" : "多元型", 
                "U" : "創意型",
                "R" : "創意型", 
                "A" : "務實型"
            },
            "protential": "音感潛能"
        },
        "L5": {
            "pattern": {
                "Ws" : "精準型", 
                "Wc" : "多元型", 
                "U" : "創意型",
                "R" : "創意型", 
                "A" : "經驗型"
            },
            "protential": "圖像潛能"
        },
        "R1": {
            "pattern": {
                "Ws" : "標竿型", 
                "Wc" : "授權型", 
                "U" : "人本型",
                "R" : "人本型", 
                "A" : "務實型"
            },
            "protential": "管理潛能"
        },
        "R2": {
            "pattern": {
                "Ws" : "線思型", 
                "Wc" : "跳躍型", 
                "U" : "整體型",
                "R" : "整體型", 
                "A" : "類比型"
            },
            "protential": "推理潛能"
        },
        "R3": {
            "pattern": {
                "Ws" : "精準型", 
                "Wc" : "多元型", 
                "U" : "創意型",
                "R" : "創意型", 
                "A" : "慣性型"
            },
            "protential": "操控潛能"
        },
        "R4": {
            "pattern": {
                "Ws" : "精準型", 
                "Wc" : "多元型", 
                "U" : "創意型",
                "R" : "創意型", 
                "A" : "務實型"
            },
            "protential": "語言潛能"
        },
        "R5": {
            "pattern": {
                "Ws" : "精準型", 
                "Wc" : "多元型", 
                "U" : "創意型",
                "R" : "創意型", 
                "A" : "務實型"
            },
            "protential": "視辨潛能"
        }
    }

    approachability_mapping = {
        "Ws": "Superior",
        "Wc": "Medium",
        "U" : "Normal",
        "R" : "Normal",
        "A" : "Normal"
    }

    persistence_mapping = {
        "Ws": "Superior",
        "A": "Superior",
        "Wc": "Medium",
        "U" : "Normal",
        "R" : "Normal"
    }

    data = {
        "report_number" : datetime.date.today().strftime("%Y%m%d"),

        #星號
        "Star_A" : "",
        "Star_U" : "",
        "Star_R" : "",
        "Star_Ws" : "",
        "Star_Wc" : "",

        #先天趨近性 & 先天堅持度
        "Approachability" : approachability_mapping[pattern_map[fingers["L1"].pattern]],
        "Persistence" : persistence_mapping[pattern_map[fingers["R1"].pattern]],

        #左手
        "L1style" : fingers["L1"].pattern,
        "L1number" : str(fingers["L1"].percentage),
        "L2style" : fingers["L2"].pattern,
        "L2number" : str(fingers["L2"].percentage),
        "L3style" : fingers["L3"].pattern,
        "L3number" : str(fingers["L3"].percentage),
        "L4style" : fingers["L4"].pattern,
        "L4number" : str(fingers["L4"].percentage),
        "L5style" : fingers["L5"].pattern,
        "L5number" : str(fingers["L5"].percentage),
        #右手
        "R1style" : fingers["R1"].pattern,
        "R1number" : str(fingers["R1"].percentage),
        "R2style" : fingers["R2"].pattern,
        "R2number" : str(fingers["R2"].percentage),
        "R3style" : fingers["R3"].pattern,
        "R3number" : str(fingers["R3"].percentage),
        "R4style" : fingers["R4"].pattern,
        "R4number" : str(fingers["R4"].percentage),
        "R5style" : fingers["R5"].pattern,
        "R5number" : str(fingers["R5"].percentage),

        "W_percentage" : str(round(pattern_percentage["W"] / 10 * 100, 0)),
        "U_percentage" : str(round(pattern_percentage["U"] / 10 * 100, 0)),
        "R_percentage" : str(round(pattern_percentage["R"] / 10 * 100, 0)),
        "A_percentage" : str(round(pattern_percentage["A"] / 10 * 100, 0)),

        #各指頭排名
        "L1_ranking" : str(fingers["L1"].ranking),
        "L2_ranking" : str(fingers["L2"].ranking),
        "L3_ranking" : str(fingers["L3"].ranking),
        "L4_ranking" : str(fingers["L4"].ranking),
        "L5_ranking" : str(fingers["L5"].ranking),
        "R1_ranking" : str(fingers["R1"].ranking),
        "R2_ranking" : str(fingers["R2"].ranking),
        "R3_ranking" : str(fingers["R3"].ranking),
        "R4_ranking" : str(fingers["R4"].ranking),
        "R5_ranking" : str(fingers["R5"].ranking),

        #各指頭加總
        "numberLR1" : str(fingers["L1"].percentage + fingers["R1"].percentage),
        "numberLR2" : str(fingers["L2"].percentage + fingers["R2"].percentage),
        "numberLR3" : str(fingers["L3"].percentage + fingers["R3"].percentage),
        "numberLR4" : str(fingers["L4"].percentage + fingers["R4"].percentage),
        "numberLR5" : str(fingers["L5"].percentage + fingers["R5"].percentage),

        "total" : str(total_value),

        #八大智慧
        "E1" : str(E["interpersonal"]),
        "E2" : str(E["introspection"]),
        "E3" : str(E["image"]),
        "E4" : str(E["math"]),
        "E5" : str(E["rhythm"]),
        "E6" : str(E["music"]),
        "E7" : str(E["language"]),
        "E8" : str(E["nature"]),

        "S1" : str((E["image"] + E["math"]) / 2),
        "S2" : str((E["image"] + E["math"] + E["nature"]) / 3),
        "S3" : str((E["nature"] + E["math"]) / 2),
        "S4" : str((E["rhythm"] + E["math"] + E["nature"]) / 3),
        "S5" : str((E["nature"] + E["math"]) / 2),
        "S6" : str((E["image"] + E["math"] + E["nature"]) / 3),
        "S7" : str((E["image"] + E["math"] + E["interpersonal"]) / 3),
        "S8" : str((E["image"] + E["rhythm"] + E["music"]) / 3),
        "S9" : str((E["interpersonal"] + E["introspection"]) / 2),
        "S10" : str((E["image"] + E["rhythm"] + E["language"]) / 3),
        "S11" : str((E["introspection"] + E["language"]) / 2),
        "S12" : str((E["interpersonal"] + E["language"]) / 2),
        "S13" : str((E["interpersonal"] + E["math"] + E["language"]) / 3),
        "S14" : str((E["interpersonal"] + E["introspection"] + E["language"]) / 3),
        "S15" : str(E["math"]),
        "S16" : str(E["rhythm"]),

        #五力分析
        "language_writing": str(f5_rankings["language"]["writing"]),
        "language_feeling": str(f5_rankings["language"]["feeling"]),
        "language_reading": str(f5_rankings["language"]["reading"]),
        "language_speaking": str(f5_rankings["language"]["speaking"]),
        "language_understanding": str(f5_rankings["language"]["understanding"]),
        "music_preformance": str(f5_rankings["music"]["preformance"]),
        "music_feeling": str(f5_rankings["music"]["feeling"]),
        "music_appreciation": str(f5_rankings["music"]["appreciation"]),
        "music_sensitivity": str(f5_rankings["music"]["sensitivity"]),
        "music_creations": str(f5_rankings["music"]["creations"]),
        "logic_calculation": str(f5_rankings["logic"]["calculation"]),
        "logic_experiment": str(f5_rankings["logic"]["experiment"]),
        "logic_induction": str(f5_rankings["logic"]["induction"]),
        "logic_deduction": str(f5_rankings["logic"]["deduction"]),
        "logic_speculation": str(f5_rankings["logic"]["speculation"]),
        "limbs_sensation": str(f5_rankings["limbs"]["sensation"]),
        "limbs_control": str(f5_rankings["limbs"]["control"]),
        "limbs_rhythm": str(f5_rankings["limbs"]["rhythm"]),
        "limbs_reaction": str(f5_rankings["limbs"]["reaction"]),
        "limbs_balance": str(f5_rankings["limbs"]["balance"]),
        "vision_assemble": str(f5_rankings["vision"]["assemble"]),
        "vision_drawing": str(f5_rankings["vision"]["drawing"]),
        "vision_image": str(f5_rankings["vision"]["image"]),
        "vision_color": str(f5_rankings["vision"]["color"]),
        "vision_space": str(f5_rankings["vision"]["space"]),
        "interpersonal_collaboration": str(f5_rankings["interpersonal"]["collaboration"]),
        "interpersonal_empathy": str(f5_rankings["interpersonal"]["empathy"]),
        "interpersonal_leadership": str(f5_rankings["interpersonal"]["leadership"]),
        "interpersonal_communication": str(f5_rankings["interpersonal"]["communication"]),
        "interpersonal_understanding": str(f5_rankings["interpersonal"]["understanding"]),
        "nature_exploration": str(f5_rankings["nature"]["exploration"]),
        "nature_concern": str(f5_rankings["nature"]["concern"]),
        "nature_identification": str(f5_rankings["nature"]["identification"]),
        "nature_observation": str(f5_rankings["nature"]["observation"]),
        "nature_categorization": str(f5_rankings["nature"]["categorization"]),
        "introspection_empathy": str(f5_rankings["introspection"]["empathy"]),
        "introspection_reflection": str(f5_rankings["introspection"]["reflection"]),
        "introspection_compassion": str(f5_rankings["introspection"]["compassion"]),
        "introspection_discipline": str(f5_rankings["introspection"]["discipline"]),
        "introspection_insight": str(f5_rankings["introspection"]["insight"]),

        #學習風格及方式
        "LearningStyle" : get_information("LearningStyle", "R2", fingers, pattern_map),
        "LearningWay" : get_information("LearningWay", "R2", fingers, pattern_map),

        #報表內頁
        "Advice" : get_information("Advice", "R2", fingers, pattern_map),
        "Strategy" : get_information("Strategy", "R2", fingers, pattern_map),

        #承壓力與消費屬性
        "Pressure" : get_information("Pressure", "R1", fingers, pattern_map),
        "Consumption" : get_information("Consumption", "R1", fingers, pattern_map),

        #工作屬性及領導風格
        "Work" : get_information("Work", "R1", fingers, pattern_map),
        "Leadership" : get_information("Leadership", "R1", fingers, pattern_map),

        #各指頭文字
        "people_L1" : get_information("L1", "L1", fingers, pattern_map),
        "people_L2" : get_information("L2", "L2", fingers, pattern_map),
        "people_L3" : get_information("L3", "L3", fingers, pattern_map),
        "people_L4" : get_information("L4", "L4", fingers, pattern_map),
        "people_L5" : get_information("L5", "L5", fingers, pattern_map),
        "people_R1" : get_information("R1", "R1", fingers, pattern_map),
        "people_R2" : get_information("R2", "R2", fingers, pattern_map),
        "people_R3" : get_information("R3", "R3", fingers, pattern_map),
        "people_R4" : get_information("R4", "R4", fingers, pattern_map),
        "people_R5" : get_information("R5", "R5", fingers, pattern_map),

        #職場發展分析
        "career_recommendation_1" : career_map[top3_fingers.get(1).name]["protential"],
        "growing_type_1" : career_map[top3_fingers.get(1).name]["pattern"][pattern_map[top3_fingers.get(1).pattern]],
        "career_detail_1" : get_information("Career", top3_fingers.get(1).name),

        "career_recommendation_2" : career_map[top3_fingers.get(2).name]["protential"],
        "growing_type_2" : career_map[top3_fingers.get(2).name]["pattern"][pattern_map[top3_fingers.get(2).pattern]],
        "career_detail_2" : get_information("Career", top3_fingers.get(2).name),

        "career_recommendation_3" : career_map[top3_fingers.get(3).name]["protential"],
        "growing_type_3" : career_map[top3_fingers.get(3).name]["pattern"][pattern_map[top3_fingers.get(3).pattern]],
        "career_detail_3" : get_information("Career", top3_fingers.get(3).name),

        "top3_career" : '、'.join([career_map[top3_fingers.get(i).name]["protential"] for i in range(1, 4)])
    }
    data[f'Star_{pattern_map[fingers["R2"].pattern]}'] = '★'

    return data

def generate_bar_chart_stream2(labels, values):
    # 创建一个 BytesIO 流来保存条形图
    bar_img_stream = io.BytesIO()

    fig, ax1 = plt.subplots(figsize=(8, 5))

    # 确保标签和值长度相同
    if len(labels) != len(values):
        raise ValueError("Labels and values must have the same length")

    colors = plt.cm.viridis(np.linspace(0.3, 0.9, len(labels)))
    bar_width = 0.5 if len(values) > 5 else 0.2
    bars = ax1.bar(labels, values, width=bar_width, color=colors, edgecolor='grey')
    ax1.set_ylim(0, max(values) + 2)  # 留出空间以显示文本标签

    # 设置刻度位置
    ax1.xaxis.set_major_locator(FixedLocator(range(len(labels))))
    rotation_degree = 45 if len(values) > 5 else 0
    ax1.set_xticklabels(labels, rotation=rotation_degree, ha="center")

    # 在每个条形旁边添加数值标签
    for bar in bars:
        yval = bar.get_height()
        ax1.annotate(f'{yval}%',
                     xy=(bar.get_x() + bar.get_width() / 2, yval),
                     xytext=(0, 3),  # 3点的垂直偏移
                     textcoords="offset points",
                     ha='center', va='bottom',
                     fontsize=9)

    def to_percent(y, position):
        return f"{int(y)}%"

    ax1.yaxis.set_major_formatter(FuncFormatter(to_percent))
    ax1.yaxis.grid(True, linestyle='--', which='major', color='grey', alpha=.25)

    # 将图表保存到 BytesIO 流中
    plt.tight_layout()
    plt.savefig(bar_img_stream, format='png')
    plt.close()

    # 将流的位置重置到开始
    bar_img_stream.seek(0)

    return bar_img_stream

def generate_bar_chart_stream(fingers):
    # 创建一个 BytesIO 流来保存条形图
    bar_img_stream = io.BytesIO()

    fig, ax1 = plt.subplots(figsize=(8, 5))

    # 条形图
    labels = [
        'Creativity', 'Imagination', 'Rhythmic', 
        'Musical', 'Visual', 'Management', 
        'Reasoning', 'Control', 'Linguistic', 
        'Observational'
    ]
    values = [finger.percentage for finger in fingers.values()]
    
    colors = plt.cm.viridis(np.linspace(0.3, 0.9, len(labels)))
    bars = ax1.bar(labels, values, color=colors, edgecolor='grey')
    ax1.set_ylim(0, max(values) + 2)  # 留出空间以显示文本标签

    # 设置刻度位置
    ax1.xaxis.set_major_locator(FixedLocator(range(len(labels))))
    ax1.set_xticklabels(labels, rotation=45, ha="right")

    # 在每个条形旁边添加数值标签
    for bar in bars:
        yval = bar.get_height()
        ax1.annotate(f'{yval}%',
                     xy=(bar.get_x() + bar.get_width() / 2, yval),
                     xytext=(0, 3),  # 3点的垂直偏移
                     textcoords="offset points",
                     ha='center', va='bottom',
                     fontsize=9)
    def to_percent(y, position):
        return f"{int(y)}%"

    ax1.yaxis.set_major_formatter(FuncFormatter(to_percent))
    ax1.yaxis.grid(True, linestyle='--', which='major', color='grey', alpha=.25)

    # 将图表保存到 BytesIO 流中
    plt.tight_layout()
    plt.savefig(bar_img_stream, format='png')
    plt.close()

    # 将流的位置重置到开始
    bar_img_stream.seek(0)

    return bar_img_stream

def generate_radar_chart_stream(fingers):
    # 创建一个 BytesIO 流来保存雷达图
    radar_img_stream = io.BytesIO()

    fig = plt.figure(figsize=(7.5, 5))
    ax2 = fig.add_subplot(111, polar=True)

    labels = [
        'Interpersonal', 'Spatial', 'Rhythmic', 
        'Musical', 'Visual', 'Management', 
        'Reasoning', 'Control', 'Linguistic', 
        'Observational'
    ]
    values = [finger.percentage for finger in fingers.values()]

    # 雷达图
    N = len(values)
    theta = np.linspace(0, 2 * np.pi, N, endpoint=False)
    values = np.concatenate((values, [values[0]]))  # Close the loop
    theta = np.concatenate((theta, [theta[0]]))  # Close the loop
    ax2.set_theta_zero_location("N")
    ax2.set_theta_direction(-1)
    ax2.plot(theta, values)
    ax2.fill(theta, values, alpha=0.25)
    ax2.set_xticks(theta[:-1])
    ax2.set_xticklabels(labels, fontsize=15)
    ax2.set_yticks(np.arange(0, 22, 2))

    # 将图表保存到 BytesIO 流中
    plt.tight_layout()
    plt.savefig(radar_img_stream, format='png')
    plt.close(fig)

    # 将流的位置重置到开始
    radar_img_stream.seek(0)

    return radar_img_stream

def replace_image_placeholders(doc, placeholders):
    for paragraph in doc.paragraphs:
        for placeholder, image_stream in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.clear()
                run = paragraph.add_run()
                image_stream.seek(0)  # 确保流的位置在开始
                run.add_picture(image_stream, width=Inches(5.8))  # 设置图片大小
                break  # 假定每个段落只有一个占位符

def customize_report(user_name, pricing_plan, fingers, pattern_map):

    data = get_data(fingers, pattern_map)

    data["name"] = user_name

    dirname = os.path.dirname(os.path.dirname(os.path.realpath(__file__)))

    # 找到相對應的模板
    template_name = pattern_map.get(fingers["R1"].pattern)
    if template_name is None:
        raise ValueError("未找到匹配的模板")
    else:
        mailmergeSource = os.path.join(dirname, "docx_template", f"{pricing_plan}_R1{template_name}.docx")

    # 使用 mailmerge 创建初始文档
    document = MailMerge(mailmergeSource)
    document.merge(**data)

    # 保存临时文档，以便于插入图表
    temp_output = os.path.join(dirname, "output", "temp_" + user_name + ".docx")
    document.write(temp_output)

    # 创建一个 Word 文档实例，以便插入图表
    doc = Document(temp_output)

    # 生成图表并替换占位符
    F10_labels = [
        'Creativity', 'Imagination', 'Rhythmic', 
        'Musical', 'Visual', 'Management', 
        'Reasoning', 'Control', 'Linguistic', 
        'Observational'
    ]
    F10_values = [finger.percentage for finger in fingers.values()]
    F10_bar_chart = generate_bar_chart_stream2(F10_labels, F10_values)

    F4_labels = [
        'Approachability', 'Persistence',
        'Aspiration', 'Dialectical'
    ]
    F4_keys = ['L1', 'R1', 'L2', 'R2']
    F4_values = [fingers[key].percentage for key in F4_keys]
    F4_bar_chart = generate_bar_chart_stream2(F4_labels, F4_values)

    radar_chart_stream = generate_radar_chart_stream(fingers)

    placehoders = {
        'IMAGE_F4_BAR_PLACEHOLDER' : F4_bar_chart, 
        'IMAGE_F10_BAR_PLACEHOLDER' : F10_bar_chart, 
        'IMAGE_RADAR_PLACEHOLDER' : radar_chart_stream
    }
    replace_image_placeholders(doc, placehoders)

    # 保存最终文档
    final_output = os.path.join(dirname, "output", f"{user_name}_{pricing_plan}.docx")
    doc.save(final_output)

    # 删除临时文件
    os.remove(temp_output)

    print(user_name + "報表已生成，並添加了圖表!")
